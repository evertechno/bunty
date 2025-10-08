"""
streamlit_converter_structured.py

Streamlit Multi-format Converter — Structured, "as-is" output.
Heuristics used:
 - PyMuPDF (fitz) to read text blocks / spans with font sizes -> detect headings vs paragraphs
 - pdfplumber to detect tabular data -> produce <table> and docx tables
 - BeautifulSoup for HTML->Text/Word conversions
 - python-docx to create Word (.docx) with headings, paragraphs, lists, tables

Notes:
 - This targets digital PDFs (embedded text). No OCR is performed.
 - Heuristics (font-size thresholds, list detection) can be tuned in the UI.
"""
import io
import os
import zipfile
import base64
from typing import List, Tuple, Dict, Any

import streamlit as st
import fitz  # PyMuPDF
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import pandas as pd
import html
import re

# -----------------------------
# Utility / Heuristic Functions
# -----------------------------

BULLET_CHARS = r"^[\u2022\u2023\u25E6\-\*\•\–\—]\s+|^\d+[\.\)]\s+"

def sanitize_text(t: str) -> str:
    return t.replace('\r', '').rstrip()

def is_bullet_line(text: str) -> bool:
    return bool(re.match(BULLET_CHARS, text.strip()))

def normalize_whitespace(s: str) -> str:
    return re.sub(r'\s+\n', '\n', re.sub(r'\n\s+', '\n', s)).strip()

def choose_heading_levels(unique_sizes: List[float]) -> Dict[float, int]:
    """
    Map font sizes -> heading levels (1..4) heuristically.
    Biggest -> h1, next -> h2, etc. If many sizes, map top ones to headings.
    """
    sizes = sorted(set(unique_sizes), reverse=True)
    mapping = {}
    # Map up to 4 distinct sizes to heading levels
    top = sizes[:4]
    for idx, s in enumerate(top):
        mapping[s] = idx + 1  # 1..4
    # any other sizes map to 0 (normal paragraph)
    return mapping

# -----------------------------
# PDF Parsing (structured)
# -----------------------------

def parse_pdf_structured(pdf_bytes: bytes, min_heading_ratio: float = 1.12) -> Dict[str, Any]:
    """
    Parse PDF into a structured intermediate representation:
    {
        "pages": [
            {
                "elements": [ {"type":"heading"/"para"/"list_item"/"table", "text":..., "level":n, ...}, ... ],
                "tables": [ { "rows":[...], "bbox":...}, ... ]
            }, ...
        ],
        "fontsizes": [list of sizes found]
    }
    Heuristics:
     - Uses PyMuPDF blocks & spans for text and font sizes.
     - Uses pdfplumber for table detection.
     - Identify bullets via regex or lines starting with bullet/number patterns.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages_out = []
    all_sizes = []
    # Gather font sizes across doc (from fits get_text("dict"))
    for p in range(len(doc)):
        page = doc.load_page(p)
        d = page.get_text("dict")
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    sz = round(span.get("size", 0), 2)
                    if sz > 0:
                        all_sizes.append(sz)

    if not all_sizes:
        unique_sizes = [12.0]
    else:
        unique_sizes = sorted(set(all_sizes), reverse=True)

    # create mapping for headings
    font_to_heading = choose_heading_levels(unique_sizes)

    # threshold to decide heading vs paragraph: if span size >= (largest * min_heading_ratio) mark as heading
    max_size = max(unique_sizes) if unique_sizes else 12.0
    heading_threshold = max_size / min_heading_ratio  # smaller denominator -> more headings

    # Now parse each page into elements
    for p in range(len(doc)):
        page = doc.load_page(p)
        d = page.get_text("dict")
        elements = []
        # first detect tables with pdfplumber on this page
        tables_page = []
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as ppdf:
                page_pl = ppdf.pages[p]
                # tables may be empty; this returns list of lists or empty
                extracted_tables = page_pl.extract_tables()
                for t in extracted_tables:
                    if t and any(any(cell for cell in row if cell not in (None, "")) for row in t):
                        tables_page.append({"rows": t})
        except Exception:
            tables_page = []

        # Add tables as elements (this keeps order approximate — we append tables at end of page parsing
        # because exact vertical ordering would require bbox comparison between pdfplumber and fitz)
        # For text blocks:
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                # skip images/other (user wanted no OCR)
                continue
            # combine lines inside a block; preserve line breaks
            block_lines = []
            for line in block.get("lines", []):
                line_text = ""
                spans = line.get("spans", [])
                # For each span, collect text; we also track max span size to infer heading
                max_span_sz = 0.0
                for span in spans:
                    stxt = span.get("text", "")
                    if stxt:
                        line_text += stxt
                    sz = span.get("size", 0)
                    if sz and sz > max_span_sz:
                        max_span_sz = sz
                if line_text.strip():
                    block_lines.append((line_text.strip(), max_span_sz))
            # Now process block_lines
            # If block contains many lines that look like bullets, mark list items
            for ln, sz in block_lines:
                ln_clean = ln.strip()
                if is_bullet_line(ln_clean):
                    # list item
                    elements.append({"type": "list_item", "text": ln_clean, "size": sz})
                else:
                    # heading detection via size or mapping
                    mapped_level = font_to_heading.get(round(sz, 2), 0)
                    if (sz >= heading_threshold) or mapped_level:
                        level = mapped_level if mapped_level else 2
                        elements.append({"type": "heading", "text": ln_clean, "level": level, "size": sz})
                    else:
                        elements.append({"type": "para", "text": ln_clean, "size": sz})
        # append detected tables (best-effort)
        for t in tables_page:
            elements.append({"type": "table", "rows": t["rows"]})

        pages_out.append({"page_number": p + 1, "elements": elements})

    return {"pages": pages_out, "fontsizes": unique_sizes}


# -----------------------------
# Converters: Intermediate -> HTML / DOCX / TEXT
# -----------------------------

def structured_to_html(parsed: dict, embed_pdf: bool = False, pdf_bytes: bytes = None) -> bytes:
    parts = ['<!doctype html><html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><style>body{font-family:Arial,Helvetica,sans-serif;line-height:1.4;padding:16px}pre{white-space:pre-wrap;}table{border-collapse:collapse;margin:8px 0}td,th{border:1px solid #ccc;padding:6px}</style></head><body>']
    for page in parsed["pages"]:
        parts.append(f'<div class="page" data-page="{page["page_number"]}" style="page-break-after:always;">')
        for el in page["elements"]:
            if el["type"] == "heading":
                lvl = min(max(int(el.get("level", 2)), 1), 6)
                text = html.escape(el["text"])
                parts.append(f"<h{lvl}>{text}</h{lvl}>")
            elif el["type"] == "para":
                text = html.escape(el["text"])
                parts.append(f"<p>{text}</p>")
            elif el["type"] == "list_item":
                # we'll build simple lists: consecutive list_items -> ul/ol
                parts.append(f"<li>{html.escape(el['text'])}</li>")
            elif el["type"] == "table":
                rows = el["rows"]
                parts.append("<table>")
                for r in rows:
                    parts.append("<tr>" + "".join(f"<td>{html.escape(str(c) if c is not None else '')}</td>" for c in r) + "</tr>")
                parts.append("</table>")
        parts.append("</div>")
    # post-process to wrap consecutive <li> into <ul>
    html_text = "\n".join(parts) + "</body></html>"
    # wrap <li> sequences into <ul>
    html_text = re.sub(r'(?:</div>\s*)?(?:\s*<li>.*?</li>\s*)+', lambda m: "<ul>" + "".join(re.findall(r'<li>.*?</li>', m.group(0))) + "</ul>", html_text, flags=re.S)
    # embed original PDF optionally
    if embed_pdf and pdf_bytes:
        b64 = base64.b64encode(pdf_bytes).decode('ascii')
        embed_snip = f'<hr/><h2>Original PDF (embedded)</h2><embed src="data:application/pdf;base64,{b64}" width="100%" height="600px"></embed>'
        html_text = html_text.replace("</body></html>", embed_snip + "</body></html>")
    return html_text.encode("utf-8")


def structured_to_text(parsed: dict) -> bytes:
    out_lines = []
    for page in parsed["pages"]:
        out_lines.append(f"--- PAGE {page['page_number']} ---")
        for el in page["elements"]:
            if el["type"] == "heading":
                out_lines.append(el["text"].upper())
            elif el["type"] == "para":
                out_lines.append(el["text"])
            elif el["type"] == "list_item":
                out_lines.append(f"- {el['text']}")
            elif el["type"] == "table":
                # simple CSV-ish representation
                rows = el["rows"]
                for r in rows:
                    out_lines.append("\t".join([str(c) if c is not None else "" for c in r]))
    joined = "\n".join(out_lines)
    return joined.encode("utf-8")


def structured_to_docx(parsed: dict) -> bytes:
    doc = Document()
    # default font size mapping
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    for page in parsed["pages"]:
        # optional page header
        doc.add_paragraph(f"--- Page {page['page_number']} ---").style = doc.styles['Normal']
        last_was_list = False
        for el in page["elements"]:
            if el["type"] == "heading":
                lvl = min(max(int(el.get("level", 2)), 1), 4)
                # use add_heading for semantic headings
                doc.add_heading(el["text"], level=lvl if lvl <= 4 else 4)
                last_was_list = False
            elif el["type"] == "para":
                doc.add_paragraph(el["text"])
                last_was_list = False
            elif el["type"] == "list_item":
                # docx doesn't have direct list API in python-docx; emulate with paragraph style 'List Bullet'
                p = doc.add_paragraph(el["text"])
                p.style = 'List Bullet'
                last_was_list = True
            elif el["type"] == "table":
                rows = el["rows"]
                if not rows:
                    continue
                # create table with n rows and n cols
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncols)
                tbl.style = 'Table Grid'
                for r in rows:
                    row_cells = tbl.add_row().cells
                    for i in range(ncols):
                        cell_text = str(r[i]) if i < len(r) and r[i] is not None else ""
                        row_cells[i].text = cell_text
                last_was_list = False
        # add a page break
        doc.add_page_break()
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------
# HTML -> Text / DOCX
# -----------------------------

def html_to_text_bytes(html_bytes: bytes) -> bytes:
    soup = BeautifulSoup(html_bytes, "html.parser")
    text = soup.get_text(separator="\n")
    return text.encode("utf-8")

def html_to_docx_bytes(html_bytes: bytes) -> bytes:
    soup = BeautifulSoup(html_bytes, "html.parser")
    doc = Document()
    for el in soup.body.descendants:
        if el.name and el.name.startswith("h") and el.get_text(strip=True):
            level = int(el.name[1]) if len(el.name) > 1 and el.name[1].isdigit() else 2
            doc.add_heading(el.get_text(strip=True), level=min(level, 4))
        elif el.name == "p" and el.get_text(strip=True):
            doc.add_paragraph(el.get_text("\n", strip=True))
        elif el.name in ("ul", "ol"):
            for li in el.find_all("li"):
                p = doc.add_paragraph(li.get_text(strip=True))
                p.style = 'List Bullet' if el.name == "ul" else 'List Number'
        elif el.name == "table":
            rows = []
            for r in el.find_all("tr"):
                cols = [c.get_text(strip=True) for c in r.find_all(["th", "td"])]
                rows.append(cols)
            if rows:
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncols)
                tbl.style = 'Table Grid'
                for r in rows:
                    cells = tbl.add_row().cells
                    for i in range(ncols):
                        cells[i].text = r[i] if i < len(r) else ""
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------
# Streamlit App UI
# -----------------------------

st.set_page_config(page_title="Legacy Converter — Preserve Structure", layout="wide", page_icon="📚")
st.title("📚 Legacy Converter — Preserve structure, create a legacy")

st.markdown("""
This app tries to **preserve structure** during conversions:
- Headings (detected via font size)
- Paragraphs (text blocks)
- Lists (bullets / numbered)
- Tables (via pdfplumber)
It uses heuristics — tune thresholds from the sidebar for better results on specific documents.
""")

with st.sidebar:
    st.header("Conversion options")
    conversion = st.selectbox("Conversion", [
        "PDF → Structured HTML",
        "PDF → Word (.docx)",
        "PDF → Plain Text",
        "HTML → Word (.docx)",
        "HTML → Plain Text"
    ])
    workers = st.number_input("Parallel workers (bulk)", min_value=1, max_value=8, value=3)
    embed_pdf = st.checkbox("Embed original PDF into HTML output (only for HTML outputs)", value=False)
    # tuning
    st.markdown("### Heuristics tuning")
    heading_ratio = st.slider("Heading size sensitivity (lower = more headings)", min_value=1.05, max_value=1.5, value=1.12, step=0.01)
    st.markdown("Tip: reduce heading sensitivity if your headings are not being detected; increase to reduce false headings.")

uploaded = st.file_uploader("Upload PDF(s) or HTML(s) — digital PDFs only (no OCR).", type=["pdf", "html"], accept_multiple_files=True)

if not uploaded:
    st.info("Upload at least one file. This converter targets digital PDFs (embedded text) and HTML files.")
    st.stop()

# Bulk processing
st.markdown(f"Files queued: {len(uploaded)}")
start = st.button("Create legacy — Convert now")

if not start:
    st.stop()

# Process files (parallelize best-effort)
from concurrent.futures import ThreadPoolExecutor, as_completed
results_for_zip = []

def process_file(uploaded_file):
    name = uploaded_file.name
    raw = uploaded_file.read()
    ext = os.path.splitext(name)[1].lower()
    try:
        if ext == ".pdf":
            parsed = parse_pdf_structured(raw, min_heading_ratio=heading_ratio)
            if conversion == "PDF → Structured HTML":
                out_bytes = structured_to_html(parsed, embed_pdf=embed_pdf, pdf_bytes=raw if embed_pdf else None)
                out_name = os.path.splitext(name)[0] + ".html"
                mime = "text/html"
            elif conversion == "PDF → Word (.docx)":
                out_bytes = structured_to_docx(parsed)
                out_name = os.path.splitext(name)[0] + ".docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif conversion == "PDF → Plain Text":
                out_bytes = structured_to_text(parsed)
                out_name = os.path.splitext(name)[0] + ".txt"
                mime = "text/plain"
            else:
                return {"name": name, "error": f"Conversion {conversion} not valid for PDF"}
        elif ext == ".html":
            if conversion == "HTML → Plain Text":
                out_bytes = html_to_text_bytes(raw)
                out_name = os.path.splitext(name)[0] + ".txt"
                mime = "text/plain"
            elif conversion == "HTML → Word (.docx)":
                out_bytes = html_to_docx_bytes(raw)
                out_name = os.path.splitext(name)[0] + ".docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                return {"name": name, "error": f"Conversion {conversion} not valid for HTML"}
        else:
            return {"name": name, "error": "Unsupported file type"}
        return {"name": name, "out_bytes": out_bytes, "out_name": out_name, "mime": mime}
    except Exception as e:
        return {"name": name, "error": str(e)}

progress = st.progress(0)
status = st.empty()
log = st.empty()

with ThreadPoolExecutor(max_workers=workers) as exe:
    futures = {exe.submit(process_file, f): f.name for f in uploaded}
    done = 0
    for fut in as_completed(futures):
        done += 1
        res = fut.result()
        progress.progress(done / len(uploaded))
        if res.get("error"):
            log.write(f"✖ {res['name']} — {res['error']}")
        else:
            results_for_zip.append(res)
            log.write(f"✔ {res['name']} → {res['out_name']} ({len(res['out_bytes']):,} bytes)")
    status.success("Conversion jobs finished")

# Show results and allow downloads; also create ZIP for bulk download
if results_for_zip:
    st.markdown("### Download converted files")
    cols = st.columns(3)
    for i, res in enumerate(results_for_zip):
        col = cols[i % 3]
        with col:
            st.write(res["out_name"])
            if res["mime"].startswith("text/"):
                preview_text = res["out_bytes"].decode("utf-8", errors="replace")[:4000]
                st.text_area(f"Preview — {res['out_name']}", preview_text, height=180)
            elif res["mime"] == "text/html":
                try:
                    st.components.v1.html(res["out_bytes"].decode("utf-8", errors="replace")[:200000], height=300, scrolling=True)
                except Exception:
                    st.write("(HTML preview failed; download instead.)")
            st.download_button("Download", data=res["out_bytes"], file_name=res["out_name"], mime=res["mime"])

    # ZIP
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for r in results_for_zip:
            zf.writestr(r["out_name"], r["out_bytes"])
    zip_buf.seek(0)
    st.download_button("Download ALL as ZIP", zip_buf.read(), file_name="converted_legacy.zip", mime="application/zip")
else:
    st.error("No successful conversions to download. Check logs above.")

st.markdown("---")
st.info("This tool preserves structure using heuristics (font sizes, bullets, table detection). If a document has odd layout or complex multi-column formatting, tweak the 'Heading size sensitivity' slider or ask me to add per-document tuning rules.")
